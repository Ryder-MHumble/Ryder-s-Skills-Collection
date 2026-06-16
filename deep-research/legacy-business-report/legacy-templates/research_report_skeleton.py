#!/usr/bin/env python3
"""
Research Report .docx Skeleton (deep-research-report v5.1)
---------------------------------------------------------
按 Ryder 偏好输出单一 .docx 调研报告（麦肯锡风）：
  - 报告标题：<主题>深度调研报告：<核心结论> (<窗口期>)
  - 文件名：<主题中文>-<核心结论简写>-<YYYY-MM-DD>.docx
  - 结构：Banner + 封面 + 目录 + 执行摘要 + 9 章正文 + 4 附录（A 检索计划 / B 证据账本 / C 信源透明度报告 / D 参考来源）
  - 视觉：麦肯锡风（2 色配色 / 0 emoji / 0 彩色 callout / 0 结尾页 / 正文内联可点击超链接）

使用方法：
  1. cp templates/research_report_skeleton.py /tmp/<project>/generate_report.py
  2. 替换 REPORT_TITLE / FILE_NAME_STEM / AS_OF_DATE 三个变量
  3. 在 9 个 # === Section N === 处填入实际研究内容
  4. 填入 4 个附录的 evidence_id / 表格数据
  5. python3 generate_report.py

依赖：python-docx 1.2+ + matplotlib（用于 banner）
  python3 -m pip install --user python-docx matplotlib
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ===================== 配置（每次替换） =====================
REPORT_TITLE = "<主题>深度调研报告"            # 必填：报告主标题
REPORT_SUBTITLE = "<核心结论>（<窗口期>）"     # 必填：副标题
REPORT_EN_SUBTITLE = "<English Subtitle>"      # 可选：英文副标题
FILE_NAME_STEM = "<主题中文>-<核心结论简写>"  # 必填：文件名主干
AS_OF_DATE = "2026-06-10"                     # 必填：YYYY-MM-DD
AUTHOR = "Ryder AI Signal · 深度调研"          # 必填：署名

# 报告元信息（封面用）
TIME_WINDOW = "2024-01-01 至 2026-06-10"      # 时间窗
GEOGRAPHY = "美国 · 英国 · 法国 · 德国"          # 地理范围
CATEGORY = "工业外骨骼（劳工保护）"            # 品类范围

# Banner 关键数据（封面右下角显示）
BANNER_KEY_METRIC = "5%"                      # 如 "5%" / "29.4%" / "$3.35B"
BANNER_METRIC_LABEL = "5%-Penetration TAM"    # 如 "5%-Penetration TAM"
BANNER_METRIC_VALUE = "$3.35B – $6.75B"       # 关键结论数值


# ===================== 麦肯锡风配色常量 =====================
COLOR_PRIMARY = RGBColor(31, 78, 121)         # 深蓝 - 标题、表头、强调
COLOR_SECONDARY = RGBColor(89, 89, 89)        # 中灰 - 副标题、辅助说明
COLOR_TERTIARY = RGBColor(127, 127, 127)      # 浅灰 - 注释、来源
COLOR_BODY = RGBColor(0, 0, 0)                # 黑色 - 正文
COLOR_ACCENT = RGBColor(192, 0, 0)            # 深红 - 仅用于关键数字
COLOR_WHITE = RGBColor(255, 255, 255)
COLOR_TABLE_ALT = 'F2F2F2'                     # 表格交替行底色
COLOR_TABLE_HEADER = '1F4E79'                  # 表头底色


# ===================== 0. Banner 图渲染 =====================

def render_banner(output_path: str) -> str:
    """用 matplotlib 渲染报告封面 banner - 深蓝科技感"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.colors import LinearSegmentedColormap
    from matplotlib.patches import FancyBboxPatch
    import numpy as np

    # 强制中文字体（macOS）
    plt.rcParams['font.sans-serif'] = ['PingFang SC', 'Heiti SC', 'STHeiti', 'Songti SC', 'sans-serif']
    plt.rcParams['axes.unicode_minus'] = False

    # 配色（与主报告一致）
    C_PRIMARY = '#1F4E79'
    C_DARK = '#0E2D4A'
    C_LIGHT = '#4A7BA8'
    C_GOLD = '#C9A961'
    C_WHITE = '#FFFFFF'

    fig, ax = plt.subplots(figsize=(9, 3), dpi=150)
    ax.set_xlim(0, 100); ax.set_ylim(0, 33)
    ax.set_aspect('equal'); ax.axis('off')

    # 背景渐变
    gradient = np.linspace(0, 1, 256).reshape(1, -1)
    cmap = LinearSegmentedColormap.from_list('mckinsey', [C_PRIMARY, C_DARK])
    ax.imshow(gradient, extent=[0, 100, 0, 33], aspect='auto', cmap=cmap, zorder=0)

    # 网格
    for x in range(0, 101, 10):
        ax.plot([x, x], [0, 33], color=C_WHITE, alpha=0.06, linewidth=0.5, zorder=1)
    for y in range(0, 34, 5):
        ax.plot([0, 100], [y, y], color=C_WHITE, alpha=0.06, linewidth=0.5, zorder=1)

    # 主数据线
    x1 = np.linspace(0, 95, 50)
    y1 = 8 + 12 * (1 - np.exp(-x1/30)) + np.random.RandomState(42).randn(50) * 0.3
    ax.plot(x1, y1, color=C_WHITE, linewidth=2.5, alpha=0.95, zorder=3)
    for x, y in [(10, 8.5), (30, 13), (50, 16), (70, 18.5), (90, 20)]:
        ax.scatter([x], [y], s=80, color=C_WHITE, zorder=4, edgecolors=C_PRIMARY, linewidths=1.5)

    # 辅线
    x2 = np.linspace(5, 92, 40)
    y2 = 5 + 8 * (1 - np.exp(-(x2-5)/25)) + np.random.RandomState(7).randn(40) * 0.4
    ax.plot(x2, y2, color=C_LIGHT, linewidth=1.5, alpha=0.7, zorder=2, linestyle='--')

    x3 = np.linspace(0, 88, 30)
    y3 = 4 + 5 * (x3/88) ** 1.5 + np.random.RandomState(13).randn(30) * 0.2
    ax.plot(x3, y3, color=C_GOLD, linewidth=1.8, alpha=0.85, zorder=2)

    # 散点
    np.random.seed(99)
    for _ in range(20):
        x = np.random.uniform(60, 95)
        y = np.random.uniform(2, 8)
        ax.scatter([x], [y], s=np.random.uniform(10, 40), color=C_WHITE, alpha=np.random.uniform(0.1, 0.4), zorder=2)

    # 顶部金色装饰条
    ax.add_patch(FancyBboxPatch((0, 31.5), 100, 0.3, boxstyle="square,pad=0", facecolor=C_GOLD, edgecolor='none', zorder=5))

    # 文字
    ax.text(2, 24, 'DEEP RESEARCH', fontsize=10, color=C_GOLD, weight='bold', zorder=6)
    ax.text(2, 20, REPORT_TITLE.upper().replace('深度调研报告', ''), fontsize=14, color=C_WHITE, weight='bold', zorder=6)
    ax.text(2, 16.5, REPORT_SUBTITLE, fontsize=10, color=C_WHITE, alpha=0.9, zorder=6)
    ax.text(2, 13, AS_OF_DATE.replace('-', '/'), fontsize=8, color=C_WHITE, alpha=0.7, zorder=6)

    # 右上品牌
    ax.text(98, 24, 'RYDER AI SIGNAL', fontsize=8, color=C_WHITE, weight='bold', alpha=0.85, ha='right', zorder=6)
    ax.text(98, 21, '深度调研 / DEEP RESEARCH', fontsize=6.5, color=C_WHITE, alpha=0.65, ha='right', zorder=6)

    # 右下关键数据
    ax.text(98, 8, BANNER_KEY_METRIC, fontsize=22, color=C_GOLD, weight='bold', ha='right', zorder=6)
    ax.text(98, 5, BANNER_METRIC_LABEL, fontsize=6.5, color=C_WHITE, alpha=0.7, ha='right', zorder=6)
    ax.text(98, 3, BANNER_METRIC_VALUE, fontsize=8, color=C_WHITE, alpha=0.9, ha='right', zorder=6, weight='bold')

    plt.tight_layout(pad=0)
    plt.savefig(output_path, dpi=150, bbox_inches='tight', pad_inches=0, facecolor=C_DARK)
    plt.close()
    return output_path


# ===================== 1. 基础工具函数 =====================

def add_hyperlink(paragraph, text, url, size=9.5):
    """在段落中添加可点击蓝色下划线超链接"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def format_table_mckinsey(table):
    """麦肯锡风格表格：白底 + 蓝头条 + 灰交替行 + 细灰边框"""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell in table.rows[0].cells:
        set_cell_shading(cell, COLOR_TABLE_HEADER)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.color.rgb = COLOR_WHITE
                r.font.bold = True
                r.font.size = Pt(10.5)
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                set_cell_shading(cell, COLOR_TABLE_ALT)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = COLOR_BODY


def make_kv_table(doc, data_rows, headers=None):
    """2 列 KV 表格（企业条目专用）"""
    if headers is None:
        headers = ['维度', '详情']
    t = doc.add_table(rows=1 + len(data_rows), cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = headers[0]
    t.rows[0].cells[1].text = headers[1]
    for i, (k, v) in enumerate(data_rows):
        t.rows[i+1].cells[0].text = str(k)
        t.rows[i+1].cells[1].text = str(v)
    format_table_mckinsey(t)
    # 加粗左列
    for row in t.rows[1:]:
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = COLOR_PRIMARY
    return t


def make_data_table(doc, headers, data_rows):
    """多列数据表格"""
    t = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    for i, row in enumerate(data_rows):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = str(val)
    format_table_mckinsey(t)
    return t


# ===================== 2. 麦肯锡风排版元素 =====================

def add_hr_line(doc):
    """细蓝色横线 - 章节分隔"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_section_cover(doc, chapter_num, chapter_title, subtitle=None):
    """麦肯锡风章节封面：横线 + 章节号 + 标题 + 副标题（可选）"""
    add_hr_line(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f'第 {chapter_num} 章')
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TERTIARY
    p2 = doc.add_paragraph()
    run = p2.add_run(chapter_title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p2.paragraph_format.space_after = Pt(2)
    if subtitle:
        p3 = doc.add_paragraph()
        run = p3.add_run(subtitle)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_SECONDARY
        run.font.italic = True
    add_hr_line(doc)
    doc.add_paragraph()


def add_h2(doc, text):
    """二级标题 - 蓝色左对齐"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_h3(doc, text):
    """三级标题 - 中灰色"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def add_para(doc, text, size=10.5, bold=False, color=None, indent=False):
    """正文段落 - 麦肯锡风格：黑色、行距 1.5"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = COLOR_BODY
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)


def add_evidence_para(doc, lead_text, body_text, evidence_id, evidence_url, evidence_label=None):
    """带内联可点击 evidence_id 的正文段落（关键：正文中 evidence_id 直接可点击）"""
    p = doc.add_paragraph()
    if lead_text:
        run = p.add_run(lead_text)
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
    if body_text:
        run = p.add_run(body_text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_BODY
    if evidence_id and evidence_url:
        run = p.add_run(' ')
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_BODY
        label = evidence_label or f'[{evidence_id}]'
        add_hyperlink(p, label, evidence_url, size=9.5)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)


def add_insight(doc, label, text):
    """关键判断/重要发现 - 麦肯锡风格（替代 ❌ 彩色 callout 块）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{label} ')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_BODY


def add_cover(doc):
    """v5.1 封面页：Banner + 标题 + 元信息"""
    # Banner 图
    banner_path = '/tmp/_report_banner.png'
    render_banner(banner_path)
    p_banner = doc.add_paragraph()
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_banner = p_banner.add_run()
    run_banner.add_picture(banner_path, width=Cm(17))

    doc.add_paragraph()

    # 主标题
    p = doc.add_paragraph()
    run = p.add_run(REPORT_TITLE)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p2 = doc.add_paragraph()
    run = p2.add_run(REPORT_SUBTITLE)
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_SECONDARY
    p2.paragraph_format.space_after = Pt(12)
    if REPORT_EN_SUBTITLE:
        p3 = doc.add_paragraph()
        run = p3.add_run(REPORT_EN_SUBTITLE)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = COLOR_TERTIARY

    for _ in range(4):
        doc.add_paragraph()

    # 元信息
    add_hr_line(doc)
    for label, val in [
        ('研究时间窗 / Time Window', TIME_WINDOW),
        ('地理范围 / Geography', GEOGRAPHY),
        ('品类范围 / Category', CATEGORY),
        ('报告版本', f'v5.1  ·  {AS_OF_DATE}'),
        ('出品方', AUTHOR),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_TERTIARY
        p2 = doc.add_paragraph()
        run = p2.add_run(val)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_BODY
    add_hr_line(doc)
    doc.add_page_break()


def add_toc(doc, items):
    """手动目录页（中英双语）"""
    p = doc.add_paragraph()
    run = p.add_run('目  录')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    add_hr_line(doc)
    doc.add_paragraph()

    for zh, en in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(zh)
        run.font.size = Pt(11.5)
        run.font.color.rgb = COLOR_PRIMARY
        run.font.bold = True
        if en:
            run2 = p.add_run(f'    /    {en}')
            run2.font.size = Pt(9.5)
            run2.font.italic = True
            run2.font.color.rgb = COLOR_TERTIARY
    doc.add_page_break()


# ===================== 3. 主流程 =====================

doc = Document()

# 中文字体设置
style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# 页面边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- 1. 封面 ---
add_cover(doc)

# --- 2. 目录 ---
add_toc(doc, [
    ('执行摘要', 'Executive Summary'),
    ('第 1 章  市场概览与驱动因素', 'Market Overview & Drivers'),
    ('第 2 章  核心需求与客户画像', 'Demand & Customer Profile'),
    ('第 3 章  主要玩家深度分析', 'Key Players Deep Dive'),
    ('第 4 章  竞争格局：波特五力', "Porter's Five Forces"),
    ('第 5 章  市场容量测算', 'Market Sizing'),
    ('第 6 章  商业模式与定价', 'Business Model & Pricing'),
    ('第 7 章  机会点与切入策略', 'Opportunities & Entry Strategy'),
    ('第 8 章  风险评估', 'Risk Assessment'),
    ('第 9 章  行动建议与下一步', 'Action Plan'),
    ('附录 A  检索计划', 'Research Plan'),
    ('附录 B  证据账本', 'Evidence Ledger'),
    ('附录 C  信源透明度报告', 'Source Transparency'),
    ('附录 D  参考来源', 'References'),
])

# --- 3. 执行摘要 ---
add_h2(doc, '执行摘要（Executive Summary）')
add_para(doc, '【TODO：引言段】')

# 5 条核心判断（用 add_insight，不要 callout）
add_h2(doc, '核心判断（5 条）')
for i in range(1, 6):
    add_insight(doc, f'判断 {i}.', '【TODO：核心结论，参考 v5.1 报告格式】')
    add_evidence_para(doc, '推论:',
        '【TODO：推论 + evidence_id 引用】',
        'E01', 'https://example.com', '[E01]')

add_h2(doc, '核心建议（一句话）')
add_para(doc, '【TODO：一句话核心建议】')

add_h2(doc, '关键数据')
# 关键数据 3 列布局
table = doc.add_table(rows=2, cols=3)
for i, label in enumerate(['数据 1', '数据 2', '数据 3']):
    table.rows[0].cells[i].text = label
for i, val in enumerate(['$0.56B', '29.4%', '$3.35B']):
    cell = table.rows[1].cells[i]
    cell.text = val
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(16)
            r.font.bold = True
            r.font.color.rgb = COLOR_PRIMARY
doc.add_page_break()

# --- 4. 正文 § 1-9 ---
for chapter_num, chapter_title, chapter_subtitle in [
    (1, '市场概览与驱动因素', 'Market Overview & Drivers'),
    (2, '核心需求与客户画像', 'Demand & Customer Profile'),
    (3, '主要玩家深度分析', 'Key Players Deep Dive'),
    (4, '竞争格局：波特五力', "Porter's Five Forces"),
    (5, '市场容量测算', 'Market Sizing'),
    (6, '商业模式与定价', 'Business Model & Pricing'),
    (7, '机会点与切入策略', 'Opportunities & Entry Strategy'),
    (8, '风险评估', 'Risk Assessment'),
    (9, '行动建议与下一步', 'Action Plan'),
]:
    add_section_cover(doc, chapter_num, chapter_title, chapter_subtitle)
    add_h2(doc, f'{chapter_num}.1 【TODO：本章第一个二级标题】')
    add_para(doc, '【TODO：叙述段 + 引用 evidence_id】')
    add_evidence_para(doc, '【TODO：关键数据】:',
        '【TODO：数据 + 说明】',
        'E01', 'https://example.com', '[E01]')
    add_insight(doc, '【TODO：关键判断/洞察】:',
        '【TODO：判断 + 推论】')
    add_para(doc, '【TODO：表格说明】')
    make_data_table(doc, ['维度', '详情'], [
        ['TODO 1', 'TODO'],
        ['TODO 2', 'TODO'],
    ])
    add_h3(doc, '本节小结')
    add_para(doc, '【TODO：本节小结 1-2 句】')
    doc.add_page_break()

# --- 5. 附录 A: 检索计划 ---
add_section_cover(doc, '附', '附录', 'Appendices')
add_h2(doc, '附录 A：检索计划 (research_plan.json)')
add_para(doc, '【TODO：说明】')
make_kv_table(doc, [
    ('topic', '【TODO】'),
    ('decision_question', '【TODO】'),
    ('as_of_date', AS_OF_DATE),
    ('time_window', TIME_WINDOW),
    ('geography', GEOGRAPHY),
    ('languages', '【TODO】'),
    ('include', '【TODO】'),
    ('exclude', '【TODO】'),
    ('source_mix', '【TODO】'),
    ('subqueries', '【TODO】'),
    ('red_lines', '【TODO】'),
    ('stop_condition', '【TODO】'),
])
doc.add_page_break()

# --- 6. 附录 B: 证据账本 ---
add_h2(doc, '附录 B：证据账本 (evidence_ledger.md)')
add_para(doc, '【TODO：说明】')
make_data_table(doc, ['E-ID', 'Claim 摘要', 'Type', 'Confidence', 'Note'], [
    ['E01', '【TODO claim】', 'B', 'high', '【TODO note】'],
    ['E02', '【TODO claim】', 'C', 'high', '【TODO note】'],
])
doc.add_page_break()

# --- 7. 附录 C: 信源透明度报告（v5.1 重写）---
add_h2(doc, '附录 C：信源透明度报告')
add_para(doc, '本附录说明报告采用的所有信源的最终采纳情况。所有信源在写作前已通过 verify_sources.py 做可达性 + 内容命中两层验证。')
add_para(doc, '· 已采纳：URL 可达 + 内容支持正文 claim，直接纳入证据账本', color=COLOR_SECONDARY)
add_para(doc, '· 备份：URL 可达但 keyword 未命中，纳入背景参考（正文未引用）', color=COLOR_SECONDARY)
add_para(doc, '· 不可达：HTTP 4xx/5xx/重定向失败，不写入报告', color=COLOR_SECONDARY)

make_data_table(doc, ['编号', '信源', '类别', '用途（对应正文 claim）', '状态'], [
    ['S-01', '【TODO 信源名】', 'B（行业研究）', '【TODO 用途】', '已采纳'],
    ['S-02', '【TODO 信源名】', 'C（公司官网）', '【TODO 用途】', '已采纳'],
    ['S-03', '【TODO 信源名】', 'C（公司官网）', '【TODO 用途】', '备份'],
    ['S-04', '【TODO 信源名】', 'C（公司官网）', '【TODO 用途】', '不可达（404）'],
])

add_para(doc, '【TODO：信源采纳统计】', color=COLOR_SECONDARY)
doc.add_page_break()

# --- 8. 附录 D: 参考来源（含超链接）---
add_h2(doc, '附录 D：参考来源 (N Evidence IDs)')
add_para(doc, '所有 evidence_id 对应的可点击超链接（按编号排序）：')

# TODO: 替换为实际 sources
sources = [
    ('E01', '【TODO 来源标题】', 'https://example.com/e01'),
    ('E02', '【TODO 来源标题】', 'https://example.com/e02'),
]
for eid, title, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{eid}. ')
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    add_hyperlink(p, title, url, size=9.5)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.8)
    p2.paragraph_format.space_after = Pt(4)
    run = p2.add_run(url)
    run.font.size = Pt(8.5)
    run.font.color.rgb = COLOR_TERTIARY

# ❌ 不加结尾页（v5.1 麦肯锡风：报告在附录 D 自然结束）

# --- 9. 保存 ---
output_dir = "/Users/rydersun/Documents/Obsidian Vault/深度调研/"
output_path = f"{output_dir}{FILE_NAME_STEM}-{AS_OF_DATE}.docx"
Path(output_dir).mkdir(parents=True, exist_ok=True)
doc.save(output_path)

size_kb = Path(output_path).stat().st_size / 1024
print(f"✅ v5.1 报告已生成: {output_path}")
print(f"   文件大小: {size_kb:.1f} KB")
print(f"   标题: {REPORT_TITLE}：{REPORT_SUBTITLE}")
print(f"   麦肯锡风：2 色配色 + 0 emoji + 0 彩色 callout + 0 结尾页")
print(f"   Banner: {BANNER_KEY_METRIC} {BANNER_METRIC_VALUE}")
